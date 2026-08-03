"""Batch process docs/needs files → mod versions with standard template.
Usage: py scripts/batch_mod.py [--limit N]
"""
import re
from pathlib import Path

NEDS_DIR = Path(__file__).resolve().parent.parent / "docs" / "needs"
OUT_DIR = NEDS_DIR  # output in same folder

# ── DOCX reader ─────────────────────────────────────────────────
try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

# ── PDF reader ──────────────────────────────────────────────────
try:
    import pypdf
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

# ── DOCX writer for mod files ───────────────────────────────────
def create_mod_docx(source_path: Path, mod_path: Path):
    """Read content from source, write formatted .mod.docx."""
    # Read source content
    text_lines = []
    title = source_path.stem

    if source_path.suffix.lower() == '.docx' and HAS_DOCX:
        try:
            doc = DocxDocument(str(source_path))
            for para in doc.paragraphs:
                t = para.text.strip()
                if t:
                    text_lines.append(t)
            if text_lines:
                title = text_lines[0][:80]
        except Exception as e:
            print(f"  WARN: {source_path.name} read failed: {e}")
            text_lines = ["[读取失败]"]
    elif source_path.suffix.lower() == '.pdf' and HAS_PDF:
        try:
            reader = pypdf.PdfReader(str(source_path))
            for page in reader.pages:
                t = page.extract_text()
                if t:
                    for line in t.split('\n'):
                        line = line.strip()
                        if line:
                            text_lines.append(line)
            if text_lines:
                title = text_lines[0][:80]
        except Exception as e:
            print(f"  WARN: {source_path.name} read failed: {e}")
            text_lines = ["[读取失败]"]
    else:
        text_lines = [f"[不支持的文件格式: {source_path.suffix}]"]

    # Write formatted mod docx
    try:
        doc = DocxDocument()

        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'SimSun'
        font.size = docx.shared.Pt(12)

        # Title
        h = doc.add_heading(title, level=1)
        doc.add_paragraph()

        # Metadata
        meta = doc.add_paragraph()
        meta_run = meta.add_run(f"来源: {source_path.name}\n处理日期: 2026-07-26\n格式: MOD 标准化版本")
        meta_run.font.size = docx.shared.Pt(9)

        doc.add_paragraph()

        # Content
        for line in text_lines:
            if not line or len(line) < 3:
                continue
            # Detect headings
            if re.match(r'^(第[一二三四五六七八九十\d]+[章节]|一[、.]|二[、.]|三[、.]|[A-Z]\d)', line) and len(line) < 60:
                doc.add_heading(line, level=2)
            elif re.match(r'^[\d]+\.[\d]*', line) and len(line) < 60:
                doc.add_heading(line, level=3)
            else:
                p = doc.add_paragraph(line)

        doc.save(str(mod_path))
        return True
    except Exception as e:
        print(f"  ERROR writing {mod_path.name}: {e}")
        return False


def main():
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--limit', type=int, default=0, help='Limit files processed')
    args = parser.parse_args()

    import docx as _docx_mod
    global docx
    docx = _docx_mod

    files = sorted(NEDS_DIR.glob('*'))

    # Separate docx, pdf, doc
    docx_files = [f for f in files if f.suffix.lower() in ('.docx', '.doc')]
    pdf_files = [f for f in files if f.suffix.lower() == '.pdf']

    if args.limit:
        docx_files = docx_files[:args.limit]

    count = 0
    total = len(docx_files) + len(pdf_files)

    print(f"Processing {total} files (DOCX: {len(docx_files)}, PDF: {len(pdf_files)})...")

    # Process DOCX files → .mod.docx
    for fp in docx_files:
        mod_path = NEDS_DIR / (fp.stem + '.mod.docx')
        print(f"[{count+1}/{total}] {fp.name} → {mod_path.name}")
        create_mod_docx(fp, mod_path)
        count += 1

    # Process PDF files → .mod.pdf (using reportlab for output)
    for fp in pdf_files:
        mod_path = NEDS_DIR / (fp.stem + '.mod.pdf')
        print(f"[{count+1}/{total}] {fp.name} → {mod_path.name}")
        try:
            # Extract text
            text_lines = []
            title = fp.stem[:80]
            if HAS_PDF:
                try:
                    reader = pypdf.PdfReader(str(fp))
                    for page in reader.pages:
                        t = page.extract_text()
                        if t:
                            for line in t.split('\n'):
                                line = line.strip()
                                if line:
                                    text_lines.append(line)
                    if text_lines:
                        title = text_lines[0][:80]
                except Exception as e:
                    text_lines = ["[读取失败]"]
                    print(f"  WARN: read failed: {e}")

            # Write formatted PDF with reportlab
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
            from reportlab.lib.units import cm
            from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

            # Register Chinese font if available
            try:
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                # Try common Windows Chinese fonts
                for font_path in [
                    r"C:\Windows\Fonts\simsun.ttc",
                    r"C:\Windows\Fonts\msyh.ttc",
                    r"C:\Windows\Fonts\msyhbd.ttc",
                ]:
                    if Path(font_path).exists():
                        pdfmetrics.registerFont(TTFont('CJK', font_path))
                        break
                else:
                    print("  WARN: No CJK font found, using Helvetica")
            except Exception:
                pass

            doc_pdf = SimpleDocTemplate(str(mod_path), pagesize=A4,
                                        rightMargin=2*cm, leftMargin=2*cm,
                                        topMargin=2*cm, bottomMargin=2*cm)
            styles = getSampleStyleSheet()
            body_style = ParagraphStyle('CJKBody', parent=styles['Normal'],
                                        fontName='CJK' if 'CJK' in locals() else 'Helvetica',
                                        fontSize=11, leading=16,
                                        spaceAfter=8)
            title_style = ParagraphStyle('CJKTitle', parent=styles['Title'],
                                         fontName='CJK' if 'CJK' in locals() else 'Helvetica',
                                         fontSize=18, leading=24,
                                         spaceAfter=16)

            story = [Paragraph(title, title_style), Spacer(1, 12)]
            for line in text_lines:
                if line and len(line) > 2:
                    story.append(Paragraph(line.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;'), body_style))

            doc_pdf.build(story)
        except Exception as e:
            print(f"  ERROR: {e}")
        count += 1

    print(f"\nDone. Processed {count}/{total} files → {NEDS_DIR}")


if __name__ == '__main__':
    main()
