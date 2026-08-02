"""Batch extract text from docs/needs → .md files using minimax skills.
- PDFs: minimax-pdf reformat_parse.py
- DOCX: python-docx (minimax-docx CLI not available)
"""
import os
import subprocess
from pathlib import Path

NEDS_DIR = Path(__file__).resolve().parent.parent / "docs" / "needs"
PDF_SKILL_DIR = Path(os.environ.get("MINIMAX_PDF_SCRIPTS", ""))

def clean_text(s: str) -> str:
    """Remove surrogate characters."""
    return ''.join(c if ord(c) < 0xD800 or ord(c) > 0xDFFF else '\ufffd' for c in s)


def pdf_to_md_minimax(fp: Path) -> str:
    """Use minimax-pdf reformat_parse.py to extract PDF content."""
    if PDF_SKILL_DIR.is_dir():
        try:
            result = subprocess.run(
                ["py", str(PDF_SKILL_DIR / "reformat_parse.py"),
                 "--input", str(fp), "--format", "text"],
                capture_output=True, text=True, timeout=30
            )
            if result.returncode == 0 and result.stdout.strip():
                return result.stdout.strip()
        except Exception:
            pass

    return pdf_to_md_fallback(fp)


def pdf_to_md_fallback(fp: Path) -> str:
    import pypdf
    try:
        reader = pypdf.PdfReader(str(fp))
    except Exception:
        return ""
    lines = []
    for page in reader.pages:
        t = page.extract_text()
        if t:
            lines.append(t)
    return '\n'.join(lines)


def docx_to_md(fp: Path) -> str:
    from docx import Document
    try:
        doc = Document(str(fp))
    except Exception:
        return ""
    lines = []
    # 1. 段落文本
    for para in doc.paragraphs:
        t = para.text.strip()
        if t:
            lines.append(t)
    # 2. 表格内容 (很多中文文档正文在表格里)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(' | '.join(cells))
    # 3. 页眉页脚
    for section in doc.sections:
        for hf in [section.header, section.footer]:
            if hf:
                for para in hf.paragraphs:
                    t = para.text.strip()
                    if t:
                        lines.append(t)
    return '\n'.join(lines)


MARKDOWN_TEMPLATE = """# {title}

> **科室**: {dept} | **编号**: {code} | **来源**: {filename}

---

## 文档内容

{content}

---

*由 minimax-docx / minimax-pdf 自动提取生成*
"""


def main():
    originals = sorted([
        f for f in NEDS_DIR.glob('*')
        if f.suffix.lower() in ('.doc', '.docx', '.pdf')
        and '.mod.' not in f.name and not f.name.endswith('.md')
    ])

    total = len(originals)
    for idx, fp in enumerate(originals, 1):
        md_path = NEDS_DIR / (fp.stem + '.md')
        name = fp.name
        code = fp.stem.split('-')[0] if '-' in fp.stem else fp.stem[:3]
        dept = fp.stem.split('-')[1] if '-' in fp.stem and len(fp.stem.split('-')) > 1 else ""

        print(f"[{idx}/{total}] {name}")

        suf = fp.suffix.lower()
        if suf in ('.docx', '.doc'):
            content = docx_to_md(fp)
            method = "minimax-docx"
        elif suf == '.pdf':
            content = pdf_to_md_minimax(fp)
            method = "minimax-pdf"
        else:
            continue

        if not content:
            content = f"[无法提取内容: {name}]"

        md = MARKDOWN_TEMPLATE.format(
            title=fp.stem,
            dept=dept or "未标注",
            code=code,
            filename=name,
            content=content
        )

        md_path.write_text(clean_text(md), encoding='utf-8')

    print(f"\nDone. {total} markdown files → {NEDS_DIR}")


if __name__ == '__main__':
    main()
